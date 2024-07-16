import os
from crewai_tools import tool, BaseTool
from pypdf import PdfReader
import functools
from docx import Document


def find_replace_in_document(doc_path, old_text, new_text):
    doc = Document(doc_path)
    
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    inline = cell.paragraphs[0].runs
                    for i in range(len(inline)):
                        if old_text in inline[i].text:
                            text = inline[i].text.replace(old_text, new_text)
                            inline[i].text = text
    
    doc.save(doc_path)

@tool("AddToScreeningMemo")
def add_to_screening_memo(section: str, data: str) -> str:
    """
    This tool is how you will fill out your screening memo. This fills in one section of the memo at a time.
    Below are the following valid sections:
        sections = [
        "<[Firm Overview]>",
        "<[Fundraise Summary & Timing]>",
        "<[Diligence Items]>",
        "<[Conclusions]>",
        "<[Overview]>",
        "<[Market Opportunity]>",
        "<[Differentiation & Winning Deals]>",
        "<[Sourcing & Picking Deals]>",
        "<[Post-Investment]>",
        "<[Exiting Deals]>",
        "<[Fund Team]>",
        "<[Co-Investment Views]>",
    ]

    """
    sections = [
        "<[Firm Overview]>",
        "<[Fundraise Summary & Timing]>",
        "<[Diligence Items]>",
        "<[Conclusions]>",
        "<[Overview]>",
        "<[Market Opportunity]>",
        "<[Differentiation & Winning Deals]>",
        "<[Sourcing & Picking Deals]>",
        "<[Post-Investment]>",
        "<[Exiting Deals]>",
        "<[Fund Team]>",
        "<[Co-Investment Views]>",
    ]
    if section not in sections:
        raise ValueError(f"The section you entered `{section}` is not valid. It must be one of the following:\n{sections}")
    out = "output.docx"
    if not os.path.exists(out):
        with open(out, 'wb') as dst, open('./templates/TEMPLATE FUND Screening Memo.docx', 'rb') as src:
            bytes = src.read()
            dst.write(bytes)
    find_replace_in_document(out, section, data)
    return "Operation completed succesfully!"   
    


class QueryDataroom(BaseTool):
    name: str = "QueryDataroom"
    description: str =     """
    Takes a filename of a file in the data room as an argument. 
    This tool will read and then return the context of that file as text.
    """

    def _run(self, filename: str) -> str:
        full_filename = f'./dataroom/{filename}'
        if not os.path.exists(full_filename):
            files = os.listdir('./dataroom')
            return f"`{filename}` does not exist in the dataroom\n \
                Files in dataroom: {files}"
        
        if full_filename.endswith('.pdf'):
            return self.read_pdf(full_filename)
        
        with open(full_filename, 'r') as f:
            return f.read()
        
    @staticmethod
    @functools.cache
    def read_pdf(filepath: str) -> str:
        reader = PdfReader(filepath)
        text = '\n\n'.join(page.extract_text() for page in reader.pages)
        return text

def create_dataroom_tool(dataroom_filepath):
    class QueryDataroom(BaseTool):
        name: str = "QueryDataroom"
        description: str =     """
        Takes a filename of a file in the data room as an argument. 
        This tool will read and then return the context of that file as text.
        """

        def _run(self, filename: str) -> str:
            full_filename = os.path.join(dataroom_filepath, filename)
            if not os.path.exists(full_filename):
                files = os.listdir(dataroom_filepath)
                return f"`{filename}` does not exist in the dataroom\n \
                    Files in dataroom: {files}"
            
            if full_filename.endswith('.pdf'):
                return self.read_pdf(full_filename)
            
            with open(full_filename, 'r') as f:
                return f.read()
            
        @staticmethod
        @functools.cache
        def read_pdf(filepath: str) -> str:
            reader = PdfReader(filepath)
            text = '\n\n'.join(page.extract_text() for page in reader.pages)
            return text
    return QueryDataroom